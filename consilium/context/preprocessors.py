"""
File preprocessors: MD / TXT / DOCX / PDF → normalized markdown with
file-name header. Each result carries its token count for fit decisions.
"""
from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Literal

import fitz
from docx import Document

from consilium.tokens import count_tokens

FileType = Literal["md", "txt", "docx", "pdf"]


class UnsupportedFileType(Exception):
    """Raised when the file extension is not in {md, txt, docx, pdf}."""


@dataclass(frozen=True)
class ProcessedFile:
    path: Path
    text: str  # normalized markdown with `# File: <name>` header
    token_count: int
    file_type: FileType


def preprocess_file(path: Path | str) -> ProcessedFile:
    """Read a file, extract its text as markdown, and count tokens."""
    path = Path(path)
    ext = path.suffix.lower().lstrip(".")

    if ext in ("md", "markdown"):
        file_type: FileType = "md"
        body = path.read_text(encoding="utf-8")
    elif ext == "txt":
        file_type = "txt"
        body = path.read_text(encoding="utf-8")
    elif ext == "docx":
        file_type = "docx"
        body = _docx_to_markdown(path)
    elif ext == "pdf":
        file_type = "pdf"
        body = _pdf_to_markdown(path)
    else:
        raise UnsupportedFileType(f"{path}: extension .{ext!r} not supported")

    text = f"# File: {path.name}\n\n{body.strip()}\n"
    return ProcessedFile(
        path=path,
        text=text,
        token_count=count_tokens(text),
        file_type=file_type,
    )


def _docx_to_markdown(path: Path) -> str:
    doc = Document(str(path))
    lines: list[str] = []
    for p in doc.paragraphs:
        style = p.style.name if p.style else ""
        if style.startswith("Heading"):
            # "Heading 1", "Heading 2", ... ; "Heading" (no number) → H1.
            suffix = style.replace("Heading", "").strip()
            try:
                level = int(suffix) if suffix else 1
            except ValueError:
                level = 1
            level = min(max(level, 1), 6)
            lines.append(f"{'#' * level} {p.text}")
        elif p.text.strip():
            lines.append(p.text)
    for tbl in doc.tables:
        if not tbl.rows:
            continue
        headers = [cell.text.strip() for cell in tbl.rows[0].cells]
        lines.append("| " + " | ".join(headers) + " |")
        lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
        for row in tbl.rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            lines.append("| " + " | ".join(cells) + " |")
        lines.append("")
    return "\n\n".join(lines)


def _pdf_to_markdown(path: Path) -> str:
    lines: list[str] = []
    with fitz.open(str(path)) as doc:
        for page_num, page in enumerate(doc, 1):
            text = page.get_text("text")
            if text.strip():
                lines.append(f"## Page {page_num}")
                lines.append(text.strip())
                lines.append("")
    return "\n".join(lines)
